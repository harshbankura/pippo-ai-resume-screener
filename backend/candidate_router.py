from fastapi import APIRouter, HTTPException, Depends, UploadFile, File, Form
from sqlalchemy.orm import Session
from sqlalchemy.exc import IntegrityError, SQLAlchemyError
from pydantic import BaseModel
from typing import Optional, List
import re
import logging
import os
import json
from datetime import datetime
from dotenv import load_dotenv
import pdfplumber
from docx import Document
from io import BytesIO
import phonenumbers
from phonenumbers import geocoder, carrier
import dateutil.parser as date_parser

from database import Candidate, SessionLocal
from ai_processing.resume_parser import enhanced_parser

load_dotenv()

logger = logging.getLogger(__name__)

router = APIRouter()

class CandidateCreate(BaseModel):
    name: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    resume_text: str
    skills: Optional[List[str]] = []
    experience_years: Optional[int] = 0
    education: Optional[str] = None
    job_description: Optional[str] = None
    job_role: Optional[str] = "Unassigned"

class CandidateResponse(CandidateCreate):
    id: int
    ai_metadata: Optional[dict] = None
    created_at: Optional[datetime] = None
    match_score: Optional[float] = None
    
    class Config:
        from_attributes = True

class CandidateUpdate(BaseModel):
    name: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    skills: Optional[List[str]] = None
    experience_years: Optional[int] = None
    education: Optional[str] = None
    match_score: Optional[float] = None
    job_role: Optional[str] = None

class BulkDeleteRequest(BaseModel):
    candidate_ids: List[int]

class SearchRequest(BaseModel):
    query: str
    filters: Optional[dict] = {}
    sort_by: Optional[str] = "match_score"
    order: Optional[str] = "desc"

def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

def validate_file_type(filename: str):
    """Validate uploaded file type"""
    if not re.search(r"\.(pdf|docx|txt)$", filename, re.IGNORECASE):
        raise HTTPException(status_code=400, detail="Only PDF, DOCX, and TXT files allowed")

def extract_text_from_file(content: bytes, filename: str) -> str:
    """Extract text from uploaded file with enhanced support"""
    try:
        if filename.lower().endswith('.pdf'):
            with pdfplumber.open(BytesIO(content)) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                return text
        elif filename.lower().endswith('.docx'):
            doc = Document(BytesIO(content))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            return text
        elif filename.lower().endswith('.txt'):
            return content.decode('utf-8', errors='ignore')
        else:
            raise ValueError("Unsupported file type")
    except Exception as e:
        logger.error(f"File extraction error: {str(e)}")
        raise HTTPException(status_code=400, detail=f"Failed to extract text from file: {str(e)}")

def validate_candidate_data(candidate_data: dict) -> dict:
    """Validate and clean candidate data"""
    # Clean name
    if candidate_data.get('name'):
        candidate_data['name'] = candidate_data['name'].strip().title()
    
    # Validate email
    if candidate_data.get('email'):
        email_pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
        if not re.match(email_pattern, candidate_data['email']):
            candidate_data['email'] = None
    
    # Validate phone
    if candidate_data.get('phone'):
        try:
            parsed_phone = phonenumbers.parse(candidate_data['phone'], "US")
            if phonenumbers.is_valid_number(parsed_phone):
                candidate_data['phone'] = phonenumbers.format_number(
                    parsed_phone, phonenumbers.PhoneNumberFormat.NATIONAL
                )
            else:
                candidate_data['phone'] = None
        except:
            candidate_data['phone'] = None
    
    # Validate experience years
    if candidate_data.get('experience_years'):
        try:
            exp_years = int(candidate_data['experience_years'])
            candidate_data['experience_years'] = max(0, min(exp_years, 50))
        except:
            candidate_data['experience_years'] = 0
    
    # Clean skills
    if candidate_data.get('skills'):
        cleaned_skills = []
        for skill in candidate_data['skills']:
            if isinstance(skill, str) and len(skill.strip()) > 1:
                cleaned_skills.append(skill.strip().title())
        candidate_data['skills'] = list(set(cleaned_skills))  # Remove duplicates
    
    return candidate_data

def calculate_overall_score(candidate_data: dict) -> int:
    """Calculate overall score (20-95 range)"""
    match_score = candidate_data.get('match_score', 0)
    experience_score = min(candidate_data.get('experience_years', 0) / 15, 1)
    skills_score = min(len(candidate_data.get('skills', [])) / 12, 1)
    contact_score = (1 if candidate_data.get('email') else 0) + (1 if candidate_data.get('phone') else 0)
    education_score = 1 if candidate_data.get('education') and candidate_data['education'].strip() else 0
    
    # Weighted calculation
    raw_score = (
        match_score * 0.35 +
        experience_score * 0.25 +
        skills_score * 0.20 +
        (contact_score / 2) * 0.10 +
        education_score * 0.10
    )
    
    # Scale to 20-95 range
    final_score = 20 + (raw_score * 75)
    
    # Bonuses and penalties
    if candidate_data.get('experience_years', 0) >= 8 and len(candidate_data.get('skills', [])) >= 15:
        final_score += 5
    
    if candidate_data.get('email') and candidate_data.get('phone') and candidate_data.get('education'):
        final_score += 3
    
    if not candidate_data.get('email') and not candidate_data.get('phone'):
        final_score -= 8
    
    return max(20, min(95, round(final_score)))

@router.post("/upload", response_model=CandidateResponse)
async def upload_resume(
    file: UploadFile = File(...),
    job_description: Optional[str] = Form(""),
    job_role: Optional[str] = Form("Unassigned"),
    db: Session = Depends(get_db)
):
    """Upload and parse resume with enhanced processing"""
    try:
        validate_file_type(file.filename)
        content = await file.read()
        
        # Extract text from file
        text = extract_text_from_file(content, file.filename)
        
        if not text.strip():
            raise HTTPException(status_code=400, detail="Could not extract text from file")
        
        # Use enhanced parser
        parsed_data = enhanced_parser.parse_text(text, job_description or "")
        
        # Validate and clean data
        parsed_data = validate_candidate_data(parsed_data)
        
        # Calculate overall score
        overall_score = calculate_overall_score(parsed_data)
        
        # Prepare AI metadata
        ai_metadata = {
            "original_filename": file.filename,
            "file_size": len(content),
            "skills_extracted": parsed_data.get("skills", []),
            "experience_years": parsed_data.get("experience_years", 0),
            "education_level": parsed_data.get("education", ""),
            "match_score": parsed_data.get("match_score", 0),
            "overall_score": overall_score,
            "job_description": job_description,
            "parsing_confidence": parsed_data.get("confidence", 0.5),
            "extraction_method": "enhanced_nlp",
            "processed_at": datetime.utcnow().isoformat()
        }
        
        # Check for duplicate email
        if parsed_data.get("email"):
            existing = db.query(Candidate).filter(Candidate.email == parsed_data["email"]).first()
            if existing:
                raise HTTPException(
                    status_code=409, 
                    detail=f"Candidate with email {parsed_data['email']} already exists"
                )
        
        # Create candidate with enhanced data
        db_candidate = Candidate(
            name=parsed_data.get("name") or "Unknown Candidate",
            email=parsed_data.get("email"),
            phone=parsed_data.get("phone"),
            resume_text=parsed_data.get("raw_text", text),
            skills=json.dumps(parsed_data.get("skills", [])),
            experience_years=parsed_data.get("experience_years", 0),
            education=parsed_data.get("education", ""),
            ai_metadata=ai_metadata,
            match_score=parsed_data.get("match_score", 0),
            job_role=job_role,
            created_at=datetime.utcnow()
        )
        
        db.add(db_candidate)
        db.commit()
        db.refresh(db_candidate)
        
        # Return response with enhanced data
        response_data = {
            "id": db_candidate.id,
            "name": db_candidate.name,
            "email": db_candidate.email,
            "phone": db_candidate.phone,
            "resume_text": db_candidate.resume_text,
            "skills": json.loads(db_candidate.skills) if db_candidate.skills else [],
            "experience_years": db_candidate.experience_years,
            "education": db_candidate.education,
            "ai_metadata": db_candidate.ai_metadata,
            "match_score": db_candidate.match_score,
            "job_role": db_candidate.job_role,
            "created_at": db_candidate.created_at
        }
        
        logger.info(f"Successfully processed resume for {db_candidate.name} (ID: {db_candidate.id})")
        return CandidateResponse(**response_data)
        
    except HTTPException as he:
        raise he
    except Exception as e:
        db.rollback()
        logger.error(f"Resume upload error: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Resume upload failed: {str(e)}")

@router.get("/candidates/", response_model=List[CandidateResponse])
def get_candidates(
    skip: int = 0, 
    limit: int = 100,
    sort_by: str = "match_score",
    order: str = "desc",
    min_experience: Optional[int] = None,
    max_experience: Optional[int] = None,
    min_match_score: Optional[float] = None,
    skills_filter: Optional[str] = None,
    db: Session = Depends(get_db)
):
    """Get candidates with advanced filtering and sorting"""
    query = db.query(Candidate)
    
    # Apply filters
    if min_experience is not None:
        query = query.filter(Candidate.experience_years >= min_experience)
    
    if max_experience is not None:
        query = query.filter(Candidate.experience_years <= max_experience)
    
    if min_match_score is not None:
        query = query.filter(Candidate.match_score >= min_match_score)
    
    if skills_filter:
        query = query.filter(Candidate.skills.ilike(f"%{skills_filter}%"))
    
    # Apply sorting
    if sort_by == "match_score":
        if order == "desc":
            query = query.order_by(Candidate.match_score.desc())
        else:
            query = query.order_by(Candidate.match_score.asc())
    elif sort_by == "created_at":
        if order == "desc":
            query = query.order_by(Candidate.created_at.desc())
        else:
            query = query.order_by(Candidate.created_at.asc())
    elif sort_by == "experience_years":
        if order == "desc":
            query = query.order_by(Candidate.experience_years.desc())
        else:
            query = query.order_by(Candidate.experience_years.asc())
    elif sort_by == "name":
        if order == "desc":
            query = query.order_by(Candidate.name.desc())
        else:
            query = query.order_by(Candidate.name.asc())
    
    candidates = query.offset(skip).limit(limit).all()
    
    result = []
    for candidate in candidates:
        candidate_data = {
            "id": candidate.id,
            "name": candidate.name,
            "email": candidate.email,
            "phone": candidate.phone,
            "resume_text": candidate.resume_text,
            "skills": json.loads(candidate.skills) if candidate.skills else [],
            "experience_years": candidate.experience_years,
            "education": candidate.education,
            "ai_metadata": candidate.ai_metadata,
            "match_score": candidate.match_score,
            "job_role": candidate.job_role,
            "created_at": candidate.created_at
        }
        result.append(CandidateResponse(**candidate_data))
    
    return result

@router.get("/candidates/{candidate_id}", response_model=CandidateResponse)
def get_candidate(candidate_id: int, db: Session = Depends(get_db)):
    """Get specific candidate by ID"""
    candidate = db.query(Candidate).filter(Candidate.id == candidate_id).first()
    if not candidate:
        raise HTTPException(status_code=404, detail="Candidate not found")
    
    candidate_data = {
        "id": candidate.id,
        "name": candidate.name,
        "email": candidate.email,
        "phone": candidate.phone,
        "resume_text": candidate.resume_text,
        "skills": json.loads(candidate.skills) if candidate.skills else [],
        "experience_years": candidate.experience_years,
        "education": candidate.education,
        "ai_metadata": candidate.ai_metadata,
        "match_score": candidate.match_score,
        "job_role": candidate.job_role,
        "created_at": candidate.created_at
    }
    
    return CandidateResponse(**candidate_data)

@router.put("/candidates/{candidate_id}", response_model=CandidateResponse)
def update_candidate(
    candidate_id: int,
    candidate_update: CandidateUpdate,
    db: Session = Depends(get_db)
):
    """Update candidate information"""
    candidate = db.query(Candidate).filter(Candidate.id == candidate_id).first()
    if not candidate:
        raise HTTPException(status_code=404, detail="Candidate not found")
    
    # Update fields
    if candidate_update.name is not None:
        candidate.name = candidate_update.name.strip().title()
    if candidate_update.email is not None:
        candidate.email = candidate_update.email
    if candidate_update.phone is not None:
        candidate.phone = candidate_update.phone
    if candidate_update.skills is not None:
        candidate.skills = json.dumps(candidate_update.skills)
    if candidate_update.experience_years is not None:
        candidate.experience_years = candidate_update.experience_years
    if candidate_update.education is not None:
        candidate.education = candidate_update.education
    if candidate_update.match_score is not None:
        candidate.match_score = candidate_update.match_score
    if candidate_update.job_role is not None:
        candidate.job_role = candidate_update.job_role
    
    # Update AI metadata
    if candidate.ai_metadata:
        candidate.ai_metadata["last_updated"] = datetime.utcnow().isoformat()
    
    db.commit()
    db.refresh(candidate)
    
    candidate_data = {
        "id": candidate.id,
        "name": candidate.name,
        "email": candidate.email,
        "phone": candidate.phone,
        "resume_text": candidate.resume_text,
        "skills": json.loads(candidate.skills) if candidate.skills else [],
        "experience_years": candidate.experience_years,
        "education": candidate.education,
        "ai_metadata": candidate.ai_metadata,
        "match_score": candidate.match_score,
        "job_role": candidate.job_role,
        "created_at": candidate.created_at
    }
    
    return CandidateResponse(**candidate_data)

@router.delete("/candidates/{candidate_id}")
def delete_candidate(candidate_id: int, db: Session = Depends(get_db)):
    """Delete specific candidate"""
    candidate = db.query(Candidate).filter(Candidate.id == candidate_id).first()
    if not candidate:
        raise HTTPException(status_code=404, detail="Candidate not found")
    
    db.delete(candidate)
    db.commit()
    
    return {"message": f"Candidate {candidate_id} deleted successfully"}

@router.post("/candidates/bulk-delete")
def bulk_delete_candidates(request: BulkDeleteRequest, db: Session = Depends(get_db)):
    """Delete multiple candidates"""
    deleted_count = db.query(Candidate).filter(Candidate.id.in_(request.candidate_ids)).delete(synchronize_session=False)
    db.commit()
    
    return {
        "message": f"Successfully deleted {deleted_count} candidates",
        "deleted_count": deleted_count
    }

@router.post("/search", response_model=List[CandidateResponse])
def search_candidates(
    request: SearchRequest,
    skip: int = 0,
    limit: int = 50,
    db: Session = Depends(get_db)
):
    """Advanced search with multiple criteria"""
    query = db.query(Candidate)
    
    # Text search across multiple fields
    if request.query:
        search_filter = db.or_(
            Candidate.name.ilike(f"%{request.query}%"),
            Candidate.email.ilike(f"%{request.query}%"),
            Candidate.skills.ilike(f"%{request.query}%"),
            Candidate.education.ilike(f"%{request.query}%")
        )
        query = query.filter(search_filter)
    
    # Apply additional filters
    if request.filters:
        if "min_experience" in request.filters:
            query = query.filter(Candidate.experience_years >= request.filters["min_experience"])
        if "max_experience" in request.filters:
            query = query.filter(Candidate.experience_years <= request.filters["max_experience"])
        if "min_match_score" in request.filters:
            query = query.filter(Candidate.match_score >= request.filters["min_match_score"])
    
    # Apply sorting
    if request.sort_by == "match_score":
        if request.order == "desc":
            query = query.order_by(Candidate.match_score.desc())
        else:
            query = query.order_by(Candidate.match_score.asc())
    elif request.sort_by == "created_at":
        if request.order == "desc":
            query = query.order_by(Candidate.created_at.desc())
        else:
            query = query.order_by(Candidate.created_at.asc())
    
    candidates = query.offset(skip).limit(limit).all()
    
    result = []
    for candidate in candidates:
        candidate_data = {
            "id": candidate.id,
            "name": candidate.name,
            "email": candidate.email,
            "phone": candidate.phone,
            "resume_text": candidate.resume_text,
            "skills": json.loads(candidate.skills) if candidate.skills else [],
            "experience_years": candidate.experience_years,
            "education": candidate.education,
            "ai_metadata": candidate.ai_metadata,
            "match_score": candidate.match_score,
            "job_role": candidate.job_role,
            "created_at": candidate.created_at
        }
        result.append(CandidateResponse(**candidate_data))
    
    return result

@router.get("/stats")
def get_stats(db: Session = Depends(get_db)):
    """Get comprehensive statistics"""
    try:
        total_candidates = db.query(Candidate).count()
        high_match_candidates = db.query(Candidate).filter(Candidate.match_score >= 0.8).count()
        experienced_candidates = db.query(Candidate).filter(Candidate.experience_years >= 5).count()
        recent_candidates = db.query(Candidate).filter(
            Candidate.created_at >= datetime.utcnow().replace(day=1)
        ).count()
        
        avg_match_score = db.query(db.func.avg(Candidate.match_score)).scalar() or 0.0
        avg_experience = db.query(db.func.avg(Candidate.experience_years)).scalar() or 0.0
        
        # Skills analysis
        all_skills = []
        candidates_with_skills = db.query(Candidate).filter(Candidate.skills.isnot(None)).all()
        for candidate in candidates_with_skills:
            try:
                skills = json.loads(candidate.skills)
                all_skills.extend(skills)
            except:
                continue
        
        skill_counts = {}
        for skill in all_skills:
            skill_counts[skill] = skill_counts.get(skill, 0) + 1
        
        top_skills = sorted(skill_counts.items(), key=lambda x: x[1], reverse=True)[:10]
        
        return {
            "total_candidates": total_candidates,
            "high_match_candidates": high_match_candidates,
            "experienced_candidates": experienced_candidates,
            "recent_candidates": recent_candidates,
            "average_match_score": round(avg_match_score, 2),
            "average_experience": round(avg_experience, 1),
            "top_skills": top_skills,
            "candidates_with_email": db.query(Candidate).filter(Candidate.email.isnot(None)).count(),
            "candidates_with_phone": db.query(Candidate).filter(Candidate.phone.isnot(None)).count(),
            "candidates_with_education": db.query(Candidate).filter(Candidate.education.isnot(None)).count()
        }
    except Exception as e:
        logger.error(f"Stats retrieval failed: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to retrieve statistics")

@router.post("/reset")
async def reset_candidates(db: Session = Depends(get_db)):
    """Reset all candidate data"""
    try:
        deleted_count = db.query(Candidate).delete(synchronize_session=False)
        db.commit()
        
        logger.info(f"Database reset: {deleted_count} candidates deleted.")
        
        return {
            "status": "success",
            "deleted": deleted_count,
            "message": f"Successfully deleted {deleted_count} candidates from the database",
            "reset_at": datetime.utcnow().isoformat()
        }
        
    except SQLAlchemyError as e:
        db.rollback()
        logger.error(f"Database reset failed: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=500, 
            detail=f"Reset failed due to database error: {str(e)}"
        )
    except Exception as e:
        db.rollback()
        logger.error(f"Unexpected error during reset: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail="An unexpected error occurred during reset"
        )

@router.get("/health")
async def health_check(db: Session = Depends(get_db)):
    """Comprehensive health check"""
    try:
        count = db.query(Candidate).count()
        latest_candidate = db.query(Candidate).order_by(Candidate.created_at.desc()).first()
        
        return {
            "status": "healthy",
            "database": "connected",
            "total_candidates": count,
            "parser_status": "enhanced_parser_active",
            "latest_upload": latest_candidate.created_at if latest_candidate else None,
            "api_version": "v1.0",
            "features": [
                "enhanced_parsing",
                "skills_extraction", 
                "experience_calculation",
                "match_scoring",
                "bulk_operations",
                "advanced_search"
            ]
        }
    except Exception as e:
        logger.error(f"Health check failed: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail="Database connection failed"
        )

@router.get("/export")
async def export_candidates(
    format: str = "json",
    include_resume_text: bool = False,
    db: Session = Depends(get_db)
):
    """Export candidates data in various formats"""
    try:
        candidates = db.query(Candidate).all()
        
        export_data = []
        for candidate in candidates:
            data = {
                "id": candidate.id,
                "name": candidate.name,
                "email": candidate.email,
                "phone": candidate.phone,
                "skills": json.loads(candidate.skills) if candidate.skills else [],
                "experience_years": candidate.experience_years,
                "education": candidate.education,
                "match_score": candidate.match_score,
                "created_at": candidate.created_at.isoformat() if candidate.created_at else None
            }
            
            if include_resume_text:
                data["resume_text"] = candidate.resume_text
            
            export_data.append(data)
        
        return {
            "format": format,
            "count": len(export_data),
            "exported_at": datetime.utcnow().isoformat(),
            "data": export_data
        }
        
    except Exception as e:
        logger.error(f"Export failed: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail="Export operation failed"
        )

@router.post("/reparse/{candidate_id}")
async def reparse_candidate(
    candidate_id: int,
    job_description: Optional[str] = Form(""),
    db: Session = Depends(get_db)
):
    """Reparse existing candidate with updated algorithms"""
    candidate = db.query(Candidate).filter(Candidate.id == candidate_id).first()
    if not candidate:
        raise HTTPException(status_code=404, detail="Candidate not found")
    
    try:
        # Reparse using enhanced parser
        parsed_data = enhanced_parser.parse_text(candidate.resume_text, job_description)
        
        # Validate and clean data
        parsed_data = validate_candidate_data(parsed_data)
        
        # Calculate new overall score
        overall_score = calculate_overall_score(parsed_data)
        
        # Update candidate
        candidate.name = parsed_data.get("name") or candidate.name
        candidate.email = parsed_data.get("email") or candidate.email
        candidate.phone = parsed_data.get("phone") or candidate.phone
        candidate.skills = json.dumps(parsed_data.get("skills", []))
        candidate.experience_years = parsed_data.get("experience_years", 0)
        candidate.education = parsed_data.get("education", "")
        candidate.match_score = parsed_data.get("match_score", 0)
        
        # Update AI metadata
        if candidate.ai_metadata:
            candidate.ai_metadata.update({
                "reparsed_at": datetime.utcnow().isoformat(),
                "overall_score": overall_score,
                "parsing_confidence": parsed_data.get("confidence", 0.5)
            })
        
        db.commit()
        db.refresh(candidate)
        
        candidate_data = {
            "id": candidate.id,
            "name": candidate.name,
            "email": candidate.email,
            "phone": candidate.phone,
            "resume_text": candidate.resume_text,
            "skills": json.loads(candidate.skills) if candidate.skills else [],
            "experience_years": candidate.experience_years,
            "education": candidate.education,
            "ai_metadata": candidate.ai_metadata,
            "match_score": candidate.match_score,
            "created_at": candidate.created_at
        }
        
        return CandidateResponse(**candidate_data)
        
    except Exception as e:
        db.rollback()
        logger.error(f"Reparse failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Reparse failed: {str(e)}")
